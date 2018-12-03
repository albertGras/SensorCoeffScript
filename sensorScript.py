import xlrd
import docx
import numpy

print("Hello World!")



#maybe use below ord function instead
#number = ord(character) - 97
FCF = 4
K1 = 5
K2 = 6
ID_RESISTOR = 4
ID_STRING = 3

#Get rid of eventually
first_coeff_loc = 13 
last_coeff_loc = 18

sensor_type = "CMF350"
print(sensor_type)
print()

#Read in sensor code file
#code_file   = open("H:\SensorScript\practice.h")
code_file   = open("C:\PVCS\ProjectsDB\Kinetis_DB\k2Src\k_src_app\coriolis\sensor.cpp")

coeff_list = ["ID String", "FCF", "K1", "GasFD", "TubeID"] #All the coeffs that need comparison
# "", "", "", "", "", "", "", "", "", "", 

#Read in ER documents
#coeff_names   = "H:\SensorScript\Sensor_Worksheet_v23.xlsx"
er_doc_blue   = "H:\SensorScript\ER docs\ER-20018334_AK.xlsx"
er_doc_red    = "H:\SensorScript\ER docs\ER-20015860_CF.xlsx"
er_doc_green  = docx.Document("H:\SensorScript\ER docs\ER-20015206_AP.docx")
r_doc_purple  = docx.Document("H:\SensorScript\ER docs\ER-20027172_AD.docx")

#Create workbooks for excel documents
#wb_coeff = xlrd.open_workbook(coeff_list)
wb_blue = xlrd.open_workbook(er_doc_blue)
wb_red = xlrd.open_workbook(er_doc_red)

#Get the correct sheet in the excel document
#coeff_sheet = wb_coeff.sheet_by_index(0)
blue_old_params = wb_blue. sheet_by_index(1)
blue_new_params = wb_blue. sheet_by_index(2)

sensor_list = [] #All the sensors listed in the code
coeff_title_list = [] 
coeff_table = [] #Table of all regular coefficients listed in the code 
cat_title_list = []
cat_table = [] #Table of all Category coefficients 
smv_title_list = []
smv_table = [] #Table of all Smart Meter Verification coefficients 

extraCharacters = ['{', ' ', '\"', '\n', 'f','}', ';', '/',]

inCoeffTable = False
inSmvTable = False
inCatTable = False

line = []
count = 0

ERdocList = []
new_row = []
beforeTables = True
title_row = False


# Read in coefficients from sensor code 
while count < 3:   # while all 3 tables havent been read (Coeff, Cat, Smv)
    line = code_file.readline()
    if not line:                  # if end of file, stop reading file
        break
    for item in extraCharacters:  #remove unwanted characters in the line
        line = line.replace(item, '')
    line = line.split(',')
    if line[0] == '':  # Remove blank lines
        continue
    if "End" in line:   # end of one of the tables
        count = count + 1
        inCoeffTable = False
        inSmvTable = False
        inCatTable = False

    elif "COEFF_TABLE" in line:      #Create coeff table list
        coeff_title_list.append(line)
        inCoeffTable = True
    elif inCoeffTable == True:       #Create coeff table
        if not any("--" in s for s in line): #Remove comment rows with additional coeff decriptions
            sensor_list.append(line[0])  # Take the sensor types from each line
            coeff_table.append(line)

    elif "CAT_TABLE" in line:        #Create cat table list
        cat_title_list.append(line)
        inCatTable = True
    elif inCatTable == True:         #Create cat table
        line[0] = line[0].split('=', 1)[-1]  #Remove all characters before the '='
        cat_table.append(line)

    elif "SMV_TABLE" in line:        #Create smv table list
        smv_title_list.append(line)
        inSmvTable = True
    elif inSmvTable == True:         #Create smv table
        line[0] = line[0].split('=', 1)[-1]  #Remove all characters before the '='
        smv_table.append(line)



#Put docx ER docs into list for easier access 
tables = er_doc_green.tables
for table in tables:
    for row in table.rows:
        title_row = False
        if(beforeTables == False and len(new_row) != 0):
            ERdocList.append(new_row)
            new_row = []
        for cell in row.cells:
            if("Sensor Model") in cell.text:
                title_row = True
                beforeTables = False
            elif(beforeTables == False and title_row == False):
                new_row.append(cell.text)
#            else:
#                print(cell.text)

# Look for sensor type in docx ER docs
for index in range(0, len(ERdocList)):
    if(sensor_type in ERdocList[index][0]):
        print("MATCH")
    print(ERdocList[index])
print()

# Populate coefficient names for title row
#for row in range(first_coeff_loc, coeff_sheet.nrows):
data = []
#new_row = []
#for row in range(first_coeff_loc, last_coeff_loc):
#    new_row.append(coeff_sheet.cell_value(row ,0))
#data.append(new_row)

#num_of_coeff = len(new_row)
#print(data)


# Finding sensor type location in ER doc
for row in range(blue_old_params.nrows):
    if sensor_type in blue_old_params.cell_value(row, 0):
        blue_old_sensor_type_loc = row
        break
for row in range(blue_new_params.nrows):
    if sensor_type in blue_new_params.cell_value(row, 0):
        blue_new_sensor_type_loc = row
        break
        

print("Old params location: ", blue_old_sensor_type_loc)
print("New params location: ", blue_new_sensor_type_loc)
print()

#Populate ER Doc row 
new_row = []
new_row.append(blue_old_params.cell_value(blue_old_sensor_type_loc, FCF))
new_row.append(blue_old_params.cell_value(blue_old_sensor_type_loc, K1))
new_row.append(blue_old_params.cell_value(blue_old_sensor_type_loc, K2))
new_row.append(blue_new_params.cell_value(blue_new_sensor_type_loc, ID_RESISTOR))
new_row.append(blue_new_params.cell_value(blue_new_sensor_type_loc, ID_STRING))
data.append(new_row)
#print(data)


# Compare ER docs and code to see if they match
# Makes the 4th row, status
#new_row = []
#for index in range(0, num_of_coeff):
#    if str(data[1][index]) == data[2][index]:
#        new_row.append("Match")
#    else:
#        new_row.append("Error")

#data.append(new_row)

#print(data[0])
#print(data[1])
#print(data[2])
#print(data[3])




print("End")












