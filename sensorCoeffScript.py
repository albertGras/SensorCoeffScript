from codeReadHelper import *
from docReadHelper import *
from excelWriteHelper import *

from docx import Document
import sys
import os

print("Hello World!")

print()


#Provide the pathes for each file and the files will be found
#inputArgs = sys.argv[1:]
#if len(inputArgs) > 1:
#    for argIndex in range(0, len(inputArgs)):
#        if os.path.exists(inputArgs[argIndex]):
#            if "20018334" in inputArgs[argIndex]:
#                blueFile = inputArgs[argIndex]
#            if "20015860" in inputArgs[argIndex]:
#                redFile = inputArgs[argIndex]
#            if "20015206" in inputArgs[argIndex]:
#                greenFile = Document(inputArgs[argIndex])
#            if "20027172" in inputArgs[argIndex]:
#                purpleFile = Document(inputArgs[argIndex])

#        else:
#            print("Error not a valid path: ", inputArgs[argIndex])



# Provde a file path and the files will be automatcially found
blueFile = 0
redFile = 0
purpleFile = 0
greenFile = 0

coeffsToCompare = 0

for arg in sys.argv[1:]:
    if os.path.isdir(arg):
        filePath = arg
        fileList = os.listdir(filePath)
        
        print("\nFiles exist")
        for file in fileList:
            if "20018334" in file and blueFile == 0:
                print(os.path.join(filePath, file))
                blueFile = os.path.join(filePath, file)

            elif "20015860" in file and redFile == 0:
                print(os.path.join(filePath, file))
                redFile = os.path.join(filePath, file)

            elif "20015206" in file and greenFile == 0:
                print(os.path.join(filePath, file))
                greenFile = Document(os.path.join(filePath, file))

            elif "20027172" in file and purpleFile == 0:
                print(os.path.join(filePath, file))
                purpleFile = Document(os.path.join(filePath, file))
                
    elif arg == "5700":
        print("\nTransmitter is 5700")
        
        #All the 5700 coeffs that need comparison
        coeffsToCompare = ["ID String", 
            "Tone Level",  "BL Temp Coeff", "Drive SP FCF", 
            "Puck P FCF",  "Max Sensor Current", 
            "Minimum Flow Multiplier", 
             "TemperatureEffect_Flow", "TemperatureEffect_Density",
            "FD Limit", "Proportional Gain 800", "PressureEffect_Density",
            "Ramp Time", 
            "dF Tone Spacing", 
            "Freq. Drift Limit",
            "NominalFlowRate", "FlowCalFactor",   "ZeroStability",  
            "I.D. Resistor", "TubeID",   "PressureEffect_Flow_Liquid", 
            "A 4", "Drive Target",  "Integral Gain 800", "Overshoot", 
            "MassFlowAccuracy_Liquid", 
            "MassFlowAccuracyMVD_Gas", "DensityAccuracy_Liquid", "Drive SP", 
            "K1", "Drive Saturation Algorithm 800", "T03", "flags",
] 

#coeffs not in 5700  ->  "Proportional Gain 2200", "Integral Gain 2200",


            
if coeffsToCompare == 0:
    print("\nIncorrect transmitter. Stopping script. Supported transmitters: 5700")
    exit()
    
if blueFile ==0 or redFile ==0 or greenFile ==0 or purpleFile == 0:
    print("\nIncorrect file location. Stopping script. Can't find ER_20018334, ER_20015860, ER_20015206, or ER_20027172")
    exit()
    
print()


codeFileLoc = os.getcwd()

supportedSensors = ["CMF010", "CMF010P", "CMF025+", "CMF050", "CMF100", "CMF200+", "CMF300+", "CMF350", "E400+IS",
                    "CMFS007", "CMFS010", "CMFS/010", "CMFS015", "CMFS/015", "CMFS/025", "CMFS025", "CMFS040", "CMFS050", 
                    "CMFS/050", "CMFS075 ", "CMFS100", "CMFS/100", "CMFS150", "CMFS/150", "F/H300", "F/H300D", "F/H300p",
                    "CMFHC2", "CMFHC3", "CMFHC4", "T025", "T050","T075", "T100", "T150", "TA010T", "TA025T", "TA050T",
                    "TA075T","TA100T","TA200T", "TA300T", "RFH/025", "R/F025HT", "RFH/050", "R/F050HT", "RFH/100", 
                    "R/F100HT", "F/R/H200", "F100P", "CNG/F050P", "HPC010P", "DH150S", "DH300S", "DS300Z", "D300", 
                    "DL200S", "D150E200",
                    "CMFS075"
                    ]

obsoleteSensors = ["A", "CMF025", "CMF300", "CMF400M",  "E400MIS", "CNG050",  "E200D150", "F025S", "F050S", 
                    "F100S", "F200S", "D006", "D012", "D025", "D040", "D065", "D100", "D150", 
                    "R/F/H025", "R/F/H050", "R/F/H100", "R/F/H200", "DS600S", 
                    "DH006S", "DH012S", "DH025S", "DH038S", "DH040S", "DH100S", "DT065H", "DT100H", "DT150H", 
                    "D012X", "DL025S", "DL025X", "DL050X", "DL065S", "DL100S", 
                    "NOTFOUND",
                    ]

# Problems in comparison spreadsheet:
#    - F/R/H200 
#    - F/H300p  
#    - D300     
#    - D150E200 

sensorToRemove = obsoleteSensors

#Lists for copying the code file
sensorList = [] #All the sensors listed in the code
mainCoeffList = [] # All of the coeffs listd in the code
coeffTable = [] #Table of all regular coefficients listed in the code 
catCoeffList = []
catTable = [] #Table of all Category coefficients 
smvCoeffList = []
smvTable = [] #Table of all Smart Meter Verification coefficients 
constantsTable = []

#Blue ER document variables
#blueFile   = "H:\ER docs\ER-20018334_AK.xlsx" #Read in blue ER document into an array
blueTitleLineNum = 15 # The row number of where the title names are located
blueDataLineNum = 21 # The row number of where the coeff data begins
newBlueCoeffList = []
newBlueTable = []
oldBlueCoeffList =[]
oldBlueTable =[]

#Red ER document variables
#redFile   = "H:\ER docs\ER-20015860_CH.xlsx" #Read in red ER document into an array
redCoeffLineNum = 6 # The row number of where the title names are located
redCoeffDataNum = 10 # The row number of where the coeff data begins
coriolisRedCoeffList = []
coriolisRedTable = []
densViscRedCoeffList = []
densViscRedTable = []


#Green ER document variables
#greenFile   = Document("H:\ER docs\ER-20015206_AP.docx") #Read in green ER document into an array
greentableOne = []
greenTableTwo = []
greenTableThree = []
greenTableFour = []

#Purple ER document variables
#purpleFile   = Document("H:\ER docs\ER-20027172_AD.docx") #Read in purple ER document into an array
purpleDocTable = []

#Lists for "compileErDocRow"
workingRow = []

#Arrays for compiling final lists
finalArray = []
finalCodeArray = []
finalDocArray = []

catTypes = []
smvTypes = []

flowLinearityTable = [] 

extraCoeffs = []
extraSensorTypes = []
sensorComparisonDict = {}


copyCodeFile(codeFileLoc, sensorList, mainCoeffList, coeffTable, catCoeffList, catTable, smvCoeffList, smvTable, constantsTable, catTypes, smvTypes)

modifiedcoeffTable = removeObsoleteSensors(coeffTable, sensorToRemove)

addCatAndSmvTablesToCoeffTable(mainCoeffList, modifiedcoeffTable, catCoeffList, catTable, smvCoeffList, smvTable, catTypes, smvTypes)

replaceVariablesWithDefinitions(modifiedcoeffTable, constantsTable)


#ER-20018334  /  Blue
copyExcelFile(oldBlueCoeffList, oldBlueTable, newBlueCoeffList, newBlueTable, blueFile, blueTitleLineNum, blueDataLineNum) 

#ER-20015860  /  Red
copyExcelFile(coriolisRedCoeffList, coriolisRedTable, densViscRedCoeffList, densViscRedTable, redFile, redCoeffLineNum, redCoeffDataNum) 

#ER-20015206 / Green
copyGreenErDoc(greenFile, greentableOne, greenTableTwo, greenTableThree, greenTableFour)

findFlowLinearityCoeffs(flowLinearityTable)


#ER-20027172 / Purple
purpleDocTable = copyPurpleErDoc(purpleFile)

createSensorComparisonDict(sensorComparisonDict, newBlueCoeffList, newBlueTable, oldBlueCoeffList, oldBlueTable)

createFinalCodeAndDocArrays(coeffsToCompare, finalCodeArray, finalDocArray, mainCoeffList, modifiedcoeffTable, newBlueCoeffList, newBlueTable, oldBlueCoeffList, oldBlueTable, coriolisRedCoeffList, coriolisRedTable, 
    densViscRedCoeffList, densViscRedTable, purpleDocTable, greentableOne, greenTableFour, sensorComparisonDict)



createFinalArray(coeffsToCompare, finalArray, finalCodeArray, finalDocArray, sensorComparisonDict)

flowLinearityMatch = compareflowLinearityTables(flowLinearityTable, greenTableTwo)

exportFinalArraytoExcelDocument(finalArray, flowLinearityTable, greenTableTwo, flowLinearityMatch)



fullDocCoeffList = newBlueCoeffList+oldBlueCoeffList+coriolisRedCoeffList+densViscRedCoeffList

fullDocSensorTypeList = []
fullDocSensorTypeList.append([i[0] for i in newBlueTable])
fullDocSensorTypeList.append([i[0] for i in oldBlueTable])
fullDocSensorTypeList.append([i[0] for i in coriolisRedTable])
fullDocSensorTypeList.append([i[0] for i in densViscRedTable])

for docCoeff in fullDocCoeffList:
    if docCoeff not in coeffsToCompare and docCoeff not in extraCoeffs:
        extraCoeffs.append(docCoeff)


for docSensorType in fullDocSensorTypeList:
    if docSensorType not in sensorList and docSensorType not in extraSensorTypes and docSensorType != '':
        extraSensorTypes.append(docSensorType)

extraCoeffs = [x for x in extraCoeffs if x != ''] #Remove blank items
extraSensorTypes = [x for x in extraSensorTypes[0] if x != ''] #Remove blank items


print()
print()
print()


print("Finished writing comparison spreadsheet")


